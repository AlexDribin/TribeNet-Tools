import tkinter as tk
from tkinter import filedialog
from tkinter.scrolledtext import ScrolledText
from docx import Document
import re
import math
from tkinter import messagebox
from tkinter import ttk

loaded_doc = None
fleet_movement_count = 0

def split_tokens(text):
    return [t.strip() for t in re.split(r"[ ,:\\]+", text) if t.strip()]

def contains_digit(token):
    return any(c.isdigit() for c in token)

def token_contains_any(token, token_list):
    return any(tok in token for tok in token_list)



def extract_truce_tokens(doc):
    tokens = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Truces"):
            parts = split_tokens(text)
            for part in parts[1:]:
                part = part.lstrip('0')
                if part:
                    tokens.append(part)
            break
    return tokens

def extract_clan_number(doc):
    for para in doc.paragraphs:
        first_line = para.text.strip()
        if first_line:
            tokens = split_tokens(first_line)
            if len(tokens) >= 2:
                raw_token = tokens[1]
                return raw_token[1:] if len(raw_token) > 1 else ""
    return ""

def is_4digit_token(token, clan_number):
    return token.isdigit() and len(token) == 4 and clan_number not in token


def print_status_lines(doc, truce_tokens, clan_number, output_box):
    any_colored_token_found = False
    output_box.insert(tk.END, "Foreigners spotted by units:\n", "highlight")
    for para in doc.paragraphs:
        text = para.text.strip()
        if "Status" in text:
            tokens = split_tokens(text)
            line_parts = []
            has_colored_token = False

            for token in tokens:
                if clan_number in token:
                    tag = "normal"
                elif token_contains_any(token, truce_tokens):
                    tag = "green"
                    has_colored_token = True
                elif contains_digit(token):
                    tag = "red"
                    has_colored_token = True
                else:
                    tag = "normal"
                line_parts.append((token + " ", tag))

            if has_colored_token:
                any_colored_token_found = True
                output_line(output_box, line_parts)

    if not any_colored_token_found:
        output_box.insert(tk.END, "No foreign units\n", "highlight")
        
def analyze_scout_sections(doc, truce_tokens, clan_number, output_box):
    sections = []
    current_section = []
    header = ""
    for para in doc.paragraphs:
        text = para.text.strip()
        if "Current Hex" in text:
            if current_section:
                sections.append((header, current_section))
                current_section = []
            header = text
        elif text.startswith("Scout"):
            current_section.append(text)
    if current_section:
        sections.append((header, current_section))

    any_matches = False  # <- Track whether anything was printed

    for header, scout_lines in sections:
        matched_lines = []
        for line in scout_lines:
            tokens = split_tokens(line)
            found = False
            for token in tokens:
                if re.fullmatch(r"\d{4}", token) and clan_number not in token:
                    found = True
                    break
            if found:
                matched_lines.append(line)

        if matched_lines:
            any_matches = True
            output_box.insert(tk.END, header + "\n", "header")
            for line in matched_lines:
                tokens = split_tokens(line)
                line_parts = []
                for token in tokens:
                    if clan_number in token:
                        tag = "normal"
                    elif token_contains_any(token, truce_tokens):
                        tag = "green"
                    elif is_4digit_token(token, clan_number):
                        tag = "red"
                    else:
                        tag = "normal"
                    line_parts.append((token + " ", tag))
                output_line(output_box, line_parts)
            output_box.insert(tk.END, "\n")

    if not any_matches:
        output_box.insert(tk.END, "No Foreign units were spotted by scouts\n", "highlight")

def output_line(box, parts):
    for text, tag in parts:
        box.insert(tk.END, text, tag)
    box.insert(tk.END, "\n")

def select_file():
    global loaded_doc
    global fleet_movement_count
    global num
    file_path = filedialog.askopenfilename(filetypes=[("Word documents", "*.docx")])
    if not file_path:
        return

    output_box.delete(1.0, tk.END)
    fleet_movement_count = 0  # Reset count for new file
    num = 0  # Reset fleet map index
    try:
        loaded_doc = Document(file_path)
        truce_tokens = extract_truce_tokens(loaded_doc)
        clan_number = extract_clan_number(loaded_doc)

        output_box.insert(tk.END, f"Report for clan: {clan_number}\n\n", "highlight")
        current_turn = print_current_turn(loaded_doc)
        output_box.insert(tk.END, current_turn, "highlight")
        print_status_lines(loaded_doc, truce_tokens, clan_number, output_box)
        analyze_scout_sections(loaded_doc, truce_tokens, clan_number, output_box)
        fleet_movement_count = find_fleet_movement_paragraphs(loaded_doc, output_box)  
  
    except Exception as e:
        output_box.insert(tk.END, f"Error: {e}", "error")

def find_fleet_movement_paragraphs(doc, output_box):
    global fleet_movement_count 
    global draw_map_button1
    global draw_map_button2
    for para in doc.paragraphs:
        words = para.text.strip().split()
        if len(words) >= 4 and words[2] == "Fleet" and words[3] == "Movement:":
            text = para.text
            move_index = text.find("Movement: Move")
            if move_index != -1:
                fleet_movement_count += 1
                move_section = text[move_index + len("Movement: Move"):].strip()
                movement_strings = move_section.split("\\")

    if fleet_movement_count == 0:
        output_box.insert(tk.END, "No fleet movement strings found.\n", "highlight")
    else:
        output_box.insert(tk.END, f"\nFleet Movements found: {fleet_movement_count}\n", "normal")
        draw_map_button2.config(state=tk.DISABLED)  # Initially disabled
        draw_map_button1.config(state=tk.NORMAL)  
    return fleet_movement_count

def print_current_turn(doc):
    """
    Extracts and prints the second paragraph's text, removing the comma and everything after it.

    Parameters:
        doc: Loaded Word document (Document object).
    """
    try:
        if len(doc.paragraphs) >= 2:
            second_para = doc.paragraphs[1].text
            cleaned = second_para.split(",")[0].strip()
            message = f"{cleaned}\n\n"
            return    message
        else:
            warning = "Document has fewer than 2 paragraphs.\n"
            return warning

    except Exception as e:
        error_msg = f"Error reading second paragraph: {e}\n"
        print(error_msg)




# Hexagon drawing constants

DIRECTIONS = {
    "N":  (0, -1),
    "NE": (1, -1),
    "SE": (1, 0),
    "S":  (0, 1),
    "SW": (-1, 1),
    "NW": (-1, 0)
}

HEX_SIZE = 40  # Radius of hexagon


def parse_direction_and_label(move_str):
    match = re.match(r"([A-Z]{1,2})[- ]([^,\\]+)", move_str.strip())
    if match:
        direction = match.group(1)
        label = match.group(2).strip()
        if direction in DIRECTIONS:
            return direction, label
    return None, None


def axial_to_pixel(q, r):
    x = HEX_SIZE * 3/2 * q
    y = HEX_SIZE * math.sqrt(3) * (r + q / 2)
    return x + 300, y + 50  # center offset


def draw_hex(canvas, x, y, label, Options = "None",size=40):
    points = []
    for i in range(6):
        angle = math.pi / 3 * i
        px = x + HEX_SIZE * math.cos(angle)
        py = y + HEX_SIZE * math.sin(angle)
        points.extend((px, py))
    
    color = get_hex_color(label)

    canvas.create_polygon(points, outline="black", fill=color, width=2)
    canvas.create_text(x, y, text=label, font=("Arial", 10, "bold"))

def get_hex_color(label):
    label = label.upper()
    if label == "START":
        return "#ffffff"
    if label == "DE":
        return "#f2e28c"  # Light Green
    if label == "AR":
        return "#f2e272"  # Light Green
    if label == "GH":
        return "#d8f296"  # Light Greenish
    if label == "BF":
        return "#d8db7f"  # Light Greenish
    if label == "BH":
        return "#d8db70"  # Light Greenish
    if label == "CH":
        return "#478920"  # Light Greenish
    if label == "D":
        return "#93c663"  # Light Greenish
    if label == "DH":
        return "#8ebc51"  # Light Greenish
    if label == "JG":
        return "#549e66"  # Light Greenish
    if label == "JH":
        return "#478451"  # Light Greenish
    if label == "RH":
        return "#cd9b00"  # Dark Yellow (Goldenrod)
    if label == "SW":
        return "#addda5"  # Light Greenish
    if label == "LCM":
        return "#567220"  # Dark Green (OliveDrab)
    if label == "LJM":
        return "#357042"  # Dark Green (ForestGreen)
    if label == "LSM":
        return "#c68e00"  # Dark Yellow (Goldenrod)
    if label == "PR":
        return "#bac97f"  # Light Yellow-Green
    if label == "O":
        return "#006699"  # Dark Blue
    if label == "L":
        return "#99ccff"  # Light Blue
    if label == "LAND":
        return "#e5e5e5"  # Light Grey
    if label == "WATER":
        return "#acb9ca"  # Light Blue-Grey

    return "white"  # Default


def advance_and_record(q, r, move_str, visited):
    """
    Parses a move string, advances axial coordinates, and records the move in visited list.
    
    Parameters:
        q (int): Current axial q coordinate.
        r (int): Current axial r coordinate.
        move_str (str): Direction-label string (e.g. "NE-PR" or "SW GH").
        visited (list): List to which (q, r, label) tuples are appended.
    
    Returns:
        tuple: Updated (q, r) coordinates.
    """
    direction, label = parse_direction_and_label(move_str)
    if direction:
        dq, dr = DIRECTIONS.get(direction, (0, 0))
        q += dq
        r += dr
        visited.append((q, r, label, "Highlight"))
    return q, r

def just_record(q, r, move_str, visited):
    """
    Same as advance_and_record but doesn't update current position—just records surrounding hexes.
    """
    direction, label = parse_direction_and_label(move_str)
    if direction:
        dq, dr = DIRECTIONS.get(direction, (0, 0))
        new_q = q + dq
        new_r = r + dr
        visited.append((new_q, new_r, label,"None"))
    return q, r  # original q, r remain unchanged

def process_extra_movements(q, r, move_str, visited):
    """
    After parsing the main move, find extra moves in brackets and call just_record up to 6 times.
    """
    extra_start = move_str.find("-(")
    if extra_start != -1:
        # Extract the substring inside the brackets
        bracket_part = move_str[extra_start + 2:].rstrip(")")
        extra_moves = [s.strip() for s in re.split(r",\s*", bracket_part) if s.strip()]

        for extra_str in extra_moves[:6]:  # Take at most 6
            q, r = just_record(q, r, extra_str, visited)
    
    return q, r
def process_sight_hexes(q, r, move_str, visited_set, visited_list):
    """
    Parses and draws 'Sight Land' and 'Sight Water' hexes after ')(' block.
    These are extra hexes in relative directions, but should not overwrite existing ones.
    """
    second_bracket = move_str.find(")(")
    if second_bracket == -1:
        return

    sight_part = move_str[second_bracket + 2:].strip().rstrip(")")

    entries = [s.strip() for s in re.split(r",\s*", sight_part) if s.strip()]
    for entry in entries[:12]:  # Only up to 12 allowed
        entry = entry.replace("Sight ", "")  # Remove "Sight"
        if "-" not in entry or "/" not in entry:
            continue

        label_part, dirs = entry.split("-", 1)
        label = label_part.strip().upper()  # LAND or WATER
        dir1_str, dir2_str = dirs.strip().split("/")

        dir1 = dir1_str.strip().upper()
        dir2 = dir2_str.strip().upper()

        # Get combined offset
        dq1, dr1 = DIRECTIONS.get(dir1, (0, 0))
        dq2, dr2 = DIRECTIONS.get(dir2, (0, 0))
        target_q = q + dq1 + dq2
        target_r = r + dr1 + dr2

        if (target_q, target_r) not in visited_set:
            visited_list.append((target_q, target_r, label,"None"))
            visited_set.add((target_q, target_r))


def draw_fleet_movement_map(canvas_content, movement_strings):
    canvas_content.delete("all")  # Clear old hexes

    q, r = 0, 0  # start at axial (0, 0)
    visited = [(q, r, "Start", "Highlight")]
    visited_set = {(q, r)}  # For fast lookup

    for move_str in movement_strings:
        q, r = advance_and_record(q, r, move_str, visited)
        visited_set.add((q, r))

        q, r = process_extra_movements(q, r, move_str, visited)
        for q1, r1, _, _ in visited:
            visited_set.add((q1, r1))

        # 🔍 New: Process sight hexes
        process_sight_hexes(q, r, move_str, visited_set, visited)


    for q, r, label, Options in visited:
        x, y = axial_to_pixel(q, r)
        draw_hex(canvas_content, x, y, label, Options)
            # 🔵 Draw a blue border around the map

    # 🔶 Draw polyline connecting highlighted hexes
    path_coords = []
    for q, r, label, highlight in visited:
        if highlight == "Highlight":
            x, y = axial_to_pixel(q, r)
            path_coords.extend((x, y))  # Collect (x, y) pairs in sequence

    if len(path_coords) >= 4:
        canvas_content.create_line(
            path_coords,
            fill="orange",  # or any color you like
            width=3,
            smooth=True  # optional: makes line slightly curved
        )


    bbox = canvas_content.bbox("all")
    if bbox:
        x1, y1, x2, y2 = bbox
        padding = 10
        canvas_content.create_rectangle(
            x1 - padding, y1 - padding, x2 + padding, y2 + padding,
            outline="blue", width=3
        )
# ✅ Update scroll region
    canvas_content.update_idletasks()
    canvas_content.configure(scrollregion=canvas_content.bbox("all"))

def draw_fleet_maps(doc, num, canvas_content):
    counter = 0
    if doc is None:
        messagebox.showerror("No File", "Please select a file first.")
        return
    # Clear previous map
    for widget in canvas_content.winfo_children():
        widget.destroy()

       # First, check for "Fleet" strings


    for para in doc.paragraphs:
        text = para.text.strip()
        if "Fleet" in text:
            for segment in text.split("\\"):
                segment = segment.strip()
                if segment.startswith("Fleet"):
                    parts = segment.split(",")
                    if len(parts) >= 3:
                        # Extract name, current_hex, and previous_hex
                        name = ",".join(parts[0:2]).strip()
                        current_hex = parts[2].strip()
                        previous_hex = ",".join(parts[3:]).strip().strip("()[]") if len(parts) > 3 else ""

            if "Fleet Movement:" in text:
                counter += 1
                if counter == num:
                    movement_start = text.find("Movement: Move")
                    if movement_start != -1:
                        move_section = text[movement_start + len("Movement: Move"):].strip()
                        movement_strings = [s.strip() for s in move_section.split("\\") if s.strip()]
                        draw_fleet_movement_map(canvas_content, movement_strings)
                        canvas_content.configure(xscrollcommand=x_scrollbar.set, yscrollcommand=y_scrollbar.set)
                        break  

def find_fleet_movement_block(doc, num, label_text, canvas_content):
#def find_fleet_movement_block(doc, num, canvas_content, x_scrollbar, y_scrollbar):
    counter = 0
    if doc is None:
        messagebox.showerror("No File", "Please select a file first.")
        return
    total_paragraphs = len(doc.paragraphs)
    i = 0  # paragraph index

    while i < total_paragraphs:
        para = doc.paragraphs[i]
        text = para.text.strip()

        # Check for "Fleet"
        fleet_found = False
        if "Fleet" in text:
            for segment in text.split("\\"):
                segment = segment.strip()
                if segment.startswith("Fleet"):
                    # Extract name, current_hex, previous_hex
                    parts = segment.split(",")
                    if len(parts) >= 3:
                        name = ",".join(parts[0:2]).strip()
                        current_hex = parts[2].strip()
                        previous_hex = ",".join(parts[3:]).strip().strip("()[]") if len(parts) > 3 else ""
                        # Optionally: print(name, current_hex, previous_hex)
                    fleet_found = True
                    break  # Only process the first "Fleet" in this paragraph

        if not fleet_found:
            i += 1
            continue  # Move to next paragraph if no "Fleet"

        # Step forward to find either "Fleet Movement" or "Status"
        i += 1
        while i < total_paragraphs:
            forward_text = doc.paragraphs[i].text.strip()
            if "Fleet Movement" in forward_text:
                counter += 1
                if counter == num:
                    # Do something (e.g., extract and draw movement)
                    movement_start = forward_text.find("Movement: Move")
                    if movement_start != -1:
                        move_section = forward_text[movement_start + len("Movement: Move"):].strip()
                        movement_strings = [s.strip() for s in move_section.split("\\") if s.strip()]
                        draw_fleet_movement_map(canvas_content, movement_strings)
                        label_text.set(name)
                        canvas_content.configure(xscrollcommand=x_scrollbar.set,
                                                 yscrollcommand=y_scrollbar.set)
                    return  # Exit after doing operation for num-th fleet movement
                else:
                    break  # Found Fleet Movement but not the right one yet → look for next Fleet
            elif "Status" in forward_text:
                break  # "Status" found → discard current Fleet and look for next one
            else:
                i += 1  # Keep moving forward if neither "Fleet Movement" nor "Status"

def fleet_map_left(loaded_doc, label_text, canvas_content):
    global draw_map_button1
    global draw_map_button2
    global num
    if num > 1:
        # Only decrement if we haven't reached the first fleet movement
        num = num - 1
        draw_map_button1.config(state=tk.NORMAL) 
        find_fleet_movement_block(loaded_doc, num, label_text, canvas_content)
        if num == 1:
            draw_map_button2.config(state=tk.DISABLED)
    
def fleet_map_right(loaded_doc, label_text, canvas_content):
    global draw_map_button2
    global draw_map_button1
    global num

    if num < fleet_movement_count:
        # Only increment if we haven't reached the last fleet movement
        num = num + 1
        if num > 1:
            draw_map_button2.config(state=tk.NORMAL) 
        find_fleet_movement_block(loaded_doc, num, label_text, canvas_content)
        if num == fleet_movement_count:
            draw_map_button1.config(state=tk.DISABLED)
         
    
# GUI Setup
root = tk.Tk()
root.title("TribeNet Turn Scanner")

# === MAIN CONTAINER FRAME ===
main_frame = tk.Frame(root)
main_frame.pack(fill=tk.BOTH, expand=True)

# === LEFT SIDE (Buttons + Output Box) ===
left_frame = tk.Frame(main_frame)
left_frame.grid(row=0, column=0, sticky="n")



# === Output Box with Scrollbar ===
output_frame = tk.Frame(left_frame)
output_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

output_scrollbar = tk.Scrollbar(output_frame)
output_scrollbar.pack(side="right", fill="y")

output_box = tk.Text(output_frame, height=30, width=80, wrap=tk.WORD, yscrollcommand=output_scrollbar.set)
output_box.pack(side="left", fill="both", padx=10, pady=10, expand=True)

output_scrollbar.config(command=output_box.yview)
# Button to select file
# Select File button (below output frame, aligned left)
file_button_frame = tk.Frame(root)
file_button_frame.pack(side=tk.TOP, anchor="w", padx=10, pady=(0, 10))

select_button = tk.Button(file_button_frame, text="Select File", command=select_file)
select_button.pack(anchor="sw", pady=(5, 2), padx=5)

# === RIGHT SIDE (Scrollable Canvas for Map) ===
right_frame = tk.Frame(main_frame)
right_frame.grid(row=0, column=1, sticky="nsew")

# Enable resizing of right_frame
main_frame.columnconfigure(1, weight=1)
main_frame.rowconfigure(0, weight=1)

# Navigation buttons frame (above canvas)
nav_frame = tk.Frame(right_frame)
nav_frame.pack(side=tk.TOP, fill=tk.X)

# Button to draw maps
num = 0
label_text = tk.StringVar()
label_text.set("Fleet")
# Middle label between Prev and Next
draw_map_button2 = tk.Button(nav_frame, text="Prev Fleet Maps", command=lambda: fleet_map_left(loaded_doc, label_text,canvas_content))
draw_map_button2.pack(side=tk.LEFT, padx=10)
draw_map_button2.config(state=tk.DISABLED)  # Initially disabled
nav_label = tk.Label(nav_frame, textvariable=label_text)
nav_label.pack(side=tk.LEFT, expand=True)


draw_map_button1 = tk.Button(nav_frame, text="Next Fleet Map", command=lambda: fleet_map_right(loaded_doc, label_text,canvas_content))
draw_map_button1.pack(side=tk.LEFT, padx=10)




# Canvas + Scrollbars
canvas_frame = tk.Frame(right_frame)
canvas_frame.pack(fill="both", expand=True)

canvas_content = tk.Canvas(canvas_frame, width=800, height=600)  # ✅ Must be a Canvas
canvas_content.pack(side="left", fill="both", expand=True)


x_scrollbar = tk.Scrollbar(canvas_frame, orient="horizontal", command=canvas_content.xview)
x_scrollbar.pack(side="bottom", fill="x")
y_scrollbar = tk.Scrollbar(canvas_frame, orient="vertical", command=canvas_content.yview)
y_scrollbar.pack(side="right", fill="y")







def on_configure(event):
    canvas_content.configure(scrollregion=canvas_content.bbox("all"))

canvas_content.bind("<Configure>", on_configure)




# Text coloring tags
output_box.tag_config("red", foreground="red")
output_box.tag_config("green", foreground="green")
output_box.tag_config("highlight", foreground="blue", font=("Arial", 10, "bold"))
output_box.tag_config("header", foreground="black", font=("Arial", 10, "bold"))
output_box.tag_config("normal", foreground="black")
output_box.tag_config("error", foreground="red", font=("Arial", 10, "bold"))

root.mainloop()
